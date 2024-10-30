# Databricks notebook source
import pandas as pd
from typing import Dict, Any, TypedDict, List
from langgraph.graph import StateGraph, END
#from IPython.display import Image, display
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import json
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import sys
import traceback
import mlflow
from langchain_databricks import ChatDatabricks
import io
import tempfile
import os
import re
from shutil import copyfileobj
#from pyspark.dbutils import DBUtils
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import gradio as gr
import base64
from databricks.sdk import WorkspaceClient
from databricks.sdk.service import catalog
import time
import random
from mlflow.metrics.genai import EvaluationExample
from transformers import GPT2Tokenizer
import mlflow
from mlflow.tracking import MlflowClient
from datetime import datetime

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

## Enable MLflow Tracing
mlflow.langchain.autolog()


# Initialize Spark session
#spark = SparkSession.builder.appName("AuditWorkflow").getOrCreate()

llm = ChatDatabricks(endpoint='databricks-meta-llama-3-1-70b-instruct')

# Set up a tokenizer
tokenizer = GPT2Tokenizer.from_pretrained("gpt2")

# Define state schema
class State(TypedDict):
    current_step: str
    raw_documents: Dict[str, Any]
    preprocessed_data: Dict[str, Any]
    risk_assessment: Dict[str, Any]
    draft_report: str
    final_report: str
    context_request: str
    errors: List[str]
    user_input: str
    user_feedback: str
    racm_data: List[Dict[str, Any]]
    agent_metrics: Dict[str, Any]
# COMMAND ----------

# Utility functions
def process_document(file_path: str, doc_type: str) -> Dict[str, Any]:
    if doc_type == "control_testing" and file_path.endswith('.xlsx'):
        return process_control_testing(file_path)
    elif file_path.endswith('.csv'):
        return pd.read_csv(file_path).to_dict()
    elif file_path.endswith('.xlsx'):
        return pd.read_excel(file_path).to_dict()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return {"raw_text": "\n".join([para.text for para in doc.paragraphs])}
    elif file_path.endswith('.json'):
        with open(file_path, 'r') as file:
            return json.load(file)
    else:
        raise ValueError(f"Unsupported file format: {file_path}")

def process_control_testing(file_path: str) -> List[Dict[str, Any]]:
    logger.info(f"Processing control testing file: {file_path}")
    df = pd.read_excel(file_path, sheet_name='RACM')
    columns = {
        'Process / Activity': 'process_activity',
        'Risk Statement': 'risk_statement',
        'Control Activities': 'control_activities',
        'Test Plan / Steps': 'test_plan_steps',
        'Test Result / Detailed Observation': 'test_result_observation',
        'Automated / Manual': 'automation_status',
        'Control Adequacy (Y/N)': 'control_adequacy'
    }
    df = df[list(columns.keys())].rename(columns=columns)
    df = df.dropna(subset=['process_activity'])
    df['automation_status'] = df['automation_status'].map({'A': 'Automatic', 'M': 'Manual'})
    df['control_adequacy'] = df['control_adequacy'].map({'Y': 'Yes', 'N': 'No'})
    # Ensure all text columns are strings and replace NaN with empty string
    for col in df.columns:
        if df[col].dtype == 'object':
            df[col] = df[col].fillna('').astype(str)
    racm_data = df.to_dict('records')
    logger.info(f"Processed {len(racm_data)} RACM items")
    return racm_data
	
# Add this function to handle MLflow experiment setup
def setup_mlflow_experiment():
    """Set up MLflow experiment for audit workflow tracking"""
    experiment_name = "audit_workflow_monitoring"
    try:
        experiment = mlflow.get_experiment_by_name(experiment_name)
        if experiment is None:
            experiment_id = mlflow.create_experiment(
                experiment_name,
                tags={"purpose": "audit_workflow_monitoring", "version": "1.0"}
            )
        else:
            experiment_id = experiment.experiment_id
        return experiment_id
    except Exception as e:
        logger.error(f"Error setting up MLflow experiment: {str(e)}")
        raise

def generate_docx_report(report_content: str, filename: str = "audit_report_23Oct.docx") -> str:
    """
    Generates a formatted audit report in DOCX format based on the given report content string.

    Args:
    - report_content: The raw audit report content string.
    - filename: The name of the DOCX file to generate.
    
    Returns:
    - The filename of the generated DOCX file.
    """
    try:
        # Initialize the document
        doc = Document()

        def add_heading(doc, text, level, centered=False):
            """ Adds a heading to the document with custom font sizes """
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            
            # Set font size and bold based on heading level
            if level == 1:
                run.font.size = Pt(18)
                run.bold = True
            elif level == 2:
                run.font.size = Pt(16)
                run.bold = True
            elif level == 3:
                run.font.size = Pt(14)
                run.bold = True
            elif level == 4:
                run.font.size = Pt(12)
                run.bold = True
            
            if centered:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def add_bullets(doc, points, level=0):
            """ Adds a list of bullet points """
            for point in points:
                p = doc.add_paragraph(point.strip(), style='List Bullet')
                p.paragraph_format.left_indent = Pt(18 * (level + 1))

        def add_table(doc, data):
            """ Adds a table """
            if len(data) < 2:  # Ensure there's at least a header and one data row
                return
            
            table = doc.add_table(rows=1, cols=len(data[0]))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, heading in enumerate(data[0]):
                hdr_cells[i].text = heading.strip()
                hdr_cells[i].paragraphs[0].runs[0].bold = True

            for row_data in data[1:]:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data.strip()

        # Process the report content
        content_parts = report_content.split("\n")
        in_stakeholder_list = False
        current_section = None
        table_data = []
        in_summary_section = False
        category_counter = 0
        issue_counter = 0
        current_category = None

        # List of standard section headers and prefixes
        standard_prefixes = ('Issue Category:', 'Observation:', 'Justification:', 'Details:', 
                            'Insights from Past Audit:', 'Recommendation:', '# Issue:', '*', '-', '+', '|')
        main_sections = ["Stakeholder List", "Executive Summary", "Summary of Audit Issues and Risk Ratings", 
                        "Key Findings and Management Actions", "Detailed Findings", "Conclusion"]

        for i, part in enumerate(content_parts):
            part = part.strip()
            if not part:
                continue

            # Clean up formatting
            part = part.replace('*', '').replace('**', '').replace('###', '').strip()

            # Handle different sections and content in the report
            if part.lower().startswith("audit title"):
                # Special handling for Audit Title
                full_title = part + " " + content_parts[i + 1].strip()
                add_heading(doc, full_title, level=1, centered=True)
                doc.add_paragraph()  # Add a new line after the title
                content_parts[i + 1] = ""  # Skip the next line
            elif part in main_sections:
                # Close the summary table section if it's open
                if in_summary_section and table_data:
                    add_table(doc, table_data)
                    table_data = []
                    in_summary_section = False
                
                # Add section heading
                add_heading(doc, part, level=1)
                current_section = part
                if part == "Stakeholder List":
                    in_stakeholder_list = True
                elif part == "Summary of Audit Issues and Risk Ratings":
                    in_summary_section = True
                elif part == "Detailed Findings":
                    category_counter = 0
                    issue_counter = 0
                    current_category = None
                else:
                    in_stakeholder_list = False
            elif part.endswith(":") and in_stakeholder_list:
                # Subsections in Stakeholder List
                add_heading(doc, part, level=3)
            elif in_stakeholder_list and (part.startswith("-") or part.startswith("+")):
                # Bullet points in Stakeholder List
                add_bullets(doc, [part[1:].strip()], level=1)
                # Check if this is the last item in the current stakeholder subsection
                if i + 1 < len(content_parts) and content_parts[i + 1].endswith(":"):
                    doc.add_paragraph()  # Add a new line after each stakeholder section
            elif part.startswith("|") and in_summary_section:
                # Table: Collect table data
                table_data.append([cell.strip() for cell in part.split("|") if cell.strip()])
            elif current_section == "Detailed Findings":
                if part.startswith('# Issue:'):
                    title = part.replace('# Issue:', '').strip()  # Don't add numbering here
                    add_heading(doc, title, level=3)
                elif not any(part.startswith(prefix) for prefix in standard_prefixes) and part not in main_sections:
                    # This must be our category header
                    category_counter += 1
                    issue_counter = 0
                    current_category = part
                    add_heading(doc, f"{category_counter}. {current_category}", level=2)
                    doc.add_paragraph()
                elif part.startswith('Issue Category:'):
                    # Skip category line as it's already in the heading
                    continue
                elif part.startswith(('Observation:', 'Justification:', 'Details:', 'Recommendation:', 'Insights from Past Audit:')):
                    # Section headers and content
                    label = part.split(':', 1)[0]
                    content = part.split(':', 1)[1].strip()
                    # Add bold section header
                    para = doc.add_paragraph()
                    run = para.add_run(f"{label}:")
                    run.bold = True
                    # Add content
                    if label == 'Recommendation':
                        # Split recommendations into bullet points
                        recommendations = [r.strip() for r in content.split('.') if r.strip()]
                        for recommendation in recommendations:
                            add_bullets(doc, [recommendation])
                    elif label == 'Insights from Past Audit':
                        # Split insights into bullet points
                        insights = [i.strip() for i in content.split('.') if i.strip()]
                        for insight in insights:
                            add_bullets(doc, [insight])
                    else:
                        doc.add_paragraph(content)
                else:
                    # Regular paragraphs
                    doc.add_paragraph(part)
            elif part.startswith("-") or part.startswith("+"):
                # Bullet points in other sections
                add_bullets(doc, [part[1:].strip()])
            else:
                # Regular paragraphs
                doc.add_paragraph(part)

        # Add the table if it hasn't been added yet (in case it's the last section)
        if table_data:
            add_table(doc, table_data)

        # Save the document
        doc.save(filename)
        return filename
    
    except Exception as e:
        logger.error(f"Error in generate_docx_report: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise  # Re-raise the exception to be caught by error handler


def process_risk_matrix(file_path: str) -> Dict[str, Any]:
    df = pd.read_excel(file_path, sheet_name='Sheet1')
    columns = {
        'Risk Rating': 'risk_rating',
        'Descriptions': 'descriptions',
        'Factors': 'factors'
    }
    df = df[list(columns.keys())].rename(columns=columns)
    df = df.dropna(subset=['risk_rating'])
    return df.to_dict('records')

def workflow_controller_agent(state: State) -> Dict[str, Any]:
    logger.info("Workflow Controller Agent executing")
    logger.info(f"Current state: {state}")
    start_time = time.time()
    token_usage = 0
    if state.get("errors"):
        logger.warning(f"Errors detected: {state['errors']}")
        return {"current_step": "error_handler", "errors": state['errors']}
    
    current_step = state.get("current_step", "preprocessor_agent")
    
    if "loop_count" not in state:
        state["loop_count"] = 0
    elif current_step == state.get("last_step"):
        state["loop_count"] += 1
    else:
        state["loop_count"] = 0
    
    if state["loop_count"] > 3:
        logger.error("Detected loop in workflow. Terminating.")
        return {"current_step": END, "errors": ["Workflow loop detected"]}
    
    state["last_step"] = current_step
    
    if not state.get("preprocessed_data"):
        next_step = "preprocessor_agent"
    elif not state.get("risk_assessment"):
        next_step = "risk_assessment_agent"
    elif not state.get("draft_report"):
        next_step = "report_generator_agent"
    else:
        next_step = END
    
    logger.info(f"Next step determined: {next_step}")
    end_time = time.time()
    elapsed_time = end_time - start_time
    token_usage = 0  # Assuming token usage is determined here

    # Update agent_metrics in the state
        
    state["agent_metrics"]["agent_name"].append("workflow_controller_agent")
    state["agent_metrics"]["total_time"].append(elapsed_time)
    state["agent_metrics"]["total_tokens"].append(token_usage)
    return {"current_step": next_step, "agent_metrics": state["agent_metrics"]}
    
def preprocessor_agent(state: State) -> Dict[str, Any]:
    logger.info("Preprocessor Agent executing")
    start_time = time.time()
    token_usage = 0
    try:
        preprocessed_data = {}
        racm_data = None
        for doc_type, file_path in state["raw_documents"].items():
            if doc_type == "control_testing":
                racm_data = process_document(file_path, doc_type)
            elif doc_type == "risk_matrix":
                preprocessed_data[doc_type] = process_risk_matrix(file_path)
            else:
                preprocessed_data[doc_type] = process_document(file_path, doc_type)
        
        end_time = time.time()
        elapsed_time = end_time - start_time
        token_usage = 0  # Assuming token usage is determined here

        # Update agent_metrics in the state
        state["agent_metrics"]["agent_name"].append("preprocessor_agent")
        state["agent_metrics"]["total_time"].append(elapsed_time)
        state["agent_metrics"]["total_tokens"].append(token_usage)
        return {
            "preprocessed_data": preprocessed_data,
            "racm_data": racm_data,
            "agent_metrics": state["agent_metrics"]
        }
        
    except Exception as e:
        logger.error(f"Error in preprocessor_agent: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return {"errors": [f"Preprocessing error: {str(e)}"], "context_request": "preprocessing_failed"}

def risk_assessment_agent(state: State) -> Dict[str, Any]:
    logger.info("Risk Assessment Agent executing")
    start_time = time.time()
    token_usage = 0
    try:
        filtered_racm_data = [
            item for item in state["racm_data"] 
            if item["control_adequacy"] == "No" 
        ]

        # Extract past audit data from preprocessed_data
        past_audit_data = state["preprocessed_data"].get("previous_audits", {}).get("raw_text", "")

        prompt = PromptTemplate(
            input_variables=["racm_data", "risk_matrix", "past_audit_data"],
            template="""
            Analyze the following RACM data, risk matrix, and past audit data:
            RACM Data: {racm_data}
            Risk Matrix: {risk_matrix}
            Past Audit Data: {past_audit_data}

            For each issue in the RACM data:
            1. Identify the title from the first line after the control reference (e.g., C.1.1, C.1.2) in the "Test Result/ Detailed Observation" of the RACM data.
            2. Collect the details of the issue until the next control reference.
            3. Provide a risk assessment rating (High-risk, Medium-risk, Low-risk) based on the Risk Matrix.
            4. Create a one or two complete sentence summary of the issue details based on the context provided in the test plan and control activities.
            5. Determine the issue category based on the "process_activity" in the RACM data.
            6. Provide a justification for the risk assessment, explaining how it corresponds to the Risk Matrix.
            7. Analyze the past audit data and extract any insights that are specifically relevant to each current issue.

            Format your response as a JSON object with the following structure:
            {{
                "assessment_rating": {{
                    "issue1": "risk_rating",
                    "issue2": "risk_rating"
                }},
                "summary": {{
                    "issue1": "summary",
                    "issue2": "summary"
                }},
                "risk_category": {{
                    "issue1": "category",
                    "issue2": "category"
                }},
                "risk_justifications": {{
                    "issue1": "justification",
                    "issue2": "justification"
                }},
                "details": {{
                    "issue1": "full_details",
                    "issue2": "full_details"
                }},
                "title": {{
                    "issue1": "title",
                    "issue2": "title"
                }},
                "past_audit_insights": {{
                    "issue1": ["relevant_insight1", "relevant_insight2"],
                    "issue2": ["relevant_insight3"]
                }}
            }}

            Ensure your response is a valid JSON object without any additional text or formatting.
            For past audit insights, only include insights that are directly relevant to each specific issue. If there are no relevant insights for an issue, provide an empty array.
            """
        )

        chain = LLMChain(llm=llm, prompt=prompt)
        
        response = chain.run(
            racm_data=json.dumps(filtered_racm_data),
            risk_matrix=json.dumps(state["preprocessed_data"]["risk_matrix"]),
            past_audit_data=past_audit_data
            )
        
        # Get the count of tokens
        tokens = tokenizer(response, return_tensors="pt")
        token_usage = len(tokens['input_ids'][0])
        
        # Log the raw response for debugging
        logger.debug(f"Raw LLM response: {response[:1000]}...")

        try:
            risk_assessment = json.loads(response)
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing JSON: {e}")
            logger.error(f"Problematic JSON: {response[:1000]}...")
            return {"errors": [f"Risk assessment error: Failed to parse risk assessment JSON. Error: {str(e)}"], 
                    "context_request": "risk_assessment_failed"}

        # Generate the output string
        output_string = "Risk Assessment Summary:\n\n"
        output_string += "\nCurrent Audit Findings:\n"
        
        for issue, title in risk_assessment['title'].items():
            output_string += f"Issue: {title}\n"
            output_string += f"Risk Rating: {risk_assessment['assessment_rating'][issue]}\n"
            output_string += f"Category: {risk_assessment['risk_category'][issue]}\n"
            output_string += f"Summary: {risk_assessment['summary'][issue]}\n"
            output_string += f"Justification: {risk_assessment['risk_justifications'][issue]}\n"
            output_string += f"Details: {risk_assessment['details'][issue]}\n"
            if issue in risk_assessment.get('past_audit_insights', {}) and risk_assessment['past_audit_insights'][issue]:
                output_string += "Past Audit Insights:\n"
                for insight in risk_assessment['past_audit_insights'][issue]:
                    output_string += f"- {insight}\n"
            output_string += "\n"

        logger.info("Risk assessment completed successfully")
        logger.debug(f"Risk assessment output: {output_string[:1000]}...")
        end_time = time.time()
        elapsed_time = end_time - start_time
        
        # Store the metrics
        state['agent_metrics']['agent_name'].append('risk_assessment_agent')
        state['agent_metrics']['total_time'].append(elapsed_time)
        state['agent_metrics']['total_tokens'].append(token_usage)

        return {"risk_assessment": risk_assessment, "risk_assessment_summary": output_string,"agent_metrics": state['agent_metrics']}
    
    except Exception as e:
        logger.error(f"Error in risk_assessment_agent: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return {"errors": [f"Risk assessment error: {str(e)}"], "context_request": "risk_assessment_failed"}
		
def report_generator_agent(state: State) -> Dict[str, Any]:
    logger.info("Report Generator Agent executing")
    start_time = time.time()
    token_usage = 0
    try:
        filtered_racm_data = [
            item for item in state["racm_data"] 
            if item["control_adequacy"] == "No" 
        ]

        prompt = PromptTemplate(
            input_variables=["risk_assessment", "racm_data"],
            template="""
            Generate a comprehensive audit report based on the following inputs:
            Risk Assessment: {risk_assessment}
            RACM Data (Issues): {racm_data}

            Follow these guidelines:
            1. Structure the report with the following sections: 

                - Audit Title
                - Stakeholder List (create a bullet list with the 4 types of stakeholders below)
                    - For Action: Then create 4 bullet list with a note for human to fill: <Write your stakeholder name here>
                    - For Information: create 4 bullet list with a note for human to fill: <Write your stakeholder name here>
                    - Audit team: create 4 bullet list with a note for human to fill: <Write your stakeholder name here>
                    - For Information: create 4 bullet list with a note for human to fill: <Write your stakeholder name here>
                - Executive Summary
                - Audit Objective:  Leave this as note for human to fill: <Write your audit objectives here>
                - Audit Entity: Leave this as note for human to fill: <Write your audit entity here>
                - Report Date: Leave this as note for human to fill: <Write your report date here>
                - Summary of Audit Issues and Risk Ratings 
                - Key Findings and Management Actions
                - Detailed Findings (with Issue in the Risk Assessment data)
                    - Issue Category: provide category in the Risk Assessment for this finding.
                    - Observation: provide summary in the Risk Assessment for this finding.
                    - Justification: provide "justification" in the Risk Assessment for this finding.
                    - Details: Provide full original text in the 'Details' part in the Risk Assessment for this finding.
                    - Insights from Past Audit: If there are relevant insights from past audits for this specific issue, list them here. If no relevant insights exist, leave a note for the human: <No insights from past audit related to this issue>
                    - Recommendation: Provide a recommendation and further actions for this finding based on the Risk Assessment, and RACM Data.
                - Conclusion

            2. Executive Summary: is audit report high-level executive summary based on the all RACM Data. No more than 5 sentences followed by bullet points mentioned above.
            3. Summary of Audit Issues and Risk Ratings: Display the Risk Assessment data in a table format. The table should have three columns with the headers: 'Issue Category', 'Observations', and 'Risk Ratings'. 
            4. Key Findings and Management Actions: summarised Current Audit key findings and actions for each observation in RACM Data. Include management agreement to implement them. 
            5. Detailed Findings: In the Current Audit Findings, ensure each issue includes all available past audit insights in a dedicated section between Details and Recommendations. With issue title, please dont use issue title number from control testing file.
            6. Conclusion: Summarize the audit report with recommendations and include high-level summary of recurring issues identified from past audits.

            Use Markdown formatting for headers and subheaders, but avoid special characters or formatting within the text itself.
            Ensure clarity and conciseness in each section, using bullet points and tables where appropriate.
            """
        )

        chain = LLMChain(llm=llm, prompt=prompt)
        response = chain.run(
            risk_assessment=json.dumps(state.get("risk_assessment", {})),
            racm_data=json.dumps(filtered_racm_data)
        )

        # Get the count of tokens
        tokens = tokenizer(response, return_tensors="pt")
        token_usage = len(tokens['input_ids'][0])

        logger.info("Generating draft DOCX report")
        docx_filename = generate_docx_report(response)
        logger.info(f"Draft DOCX report generated: {docx_filename}")
        end_time = time.time()
        elapsed_time = end_time - start_time
        

        # Update agent_metrics in the state
        
        
        state["agent_metrics"]["agent_name"].append("report_generator_agent")
        state["agent_metrics"]["total_time"].append(elapsed_time)
        state["agent_metrics"]["total_tokens"].append(token_usage)
        return {
            "draft_report": response,
            "docx_report": docx_filename,
            "current_step": "workflow_controller_agent",
            "agent_metrics": state["agent_metrics"]
        }
    except Exception as e:
        logger.error(f"Error in report_generator_agent: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return {
            "errors": [f"Report generation error: {str(e)}"],
            "context_request": "report_generation_failed",
            "current_step": "workflow_controller_agent"
        }

def error_handler(state: State) -> Dict[str, Any]:
    logger.info("Error Handler executing")
    errors = state.get("errors", [])
    logger.warning(f"Handling errors: {errors}")
    
    try:
        if "retry_count" not in state:
            state["retry_count"] = 0
        
        if state["retry_count"] < 3:  # Allow up to 3 retries
            state["retry_count"] += 1
            logger.info(f"Retrying... Attempt {state['retry_count']}")
            return {"errors": [], "current_step": state.get("last_successful_step", "workflow_controller_agent")}
        else:
            logger.error("Max retries reached. Terminating workflow.")
            return {"current_step": END, "final_report": f"Audit terminated due to persistent errors: {errors}"}
    except Exception as e:
        logger.error(f"Error in error_handler: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return {"current_step": END, "final_report": f"Error handler failed: {str(e)}"}

def user_interaction_handler(state: State) -> Dict[str, Any]:
    logger.info("User Interaction Handler executing")
    start_time = time.time()
    token_usage = 0 
    user_input = state.get("user_input", "")
    try:
        if "help" in user_input.lower():
            response = "Available commands: 'status' to check workflow progress, 'pause' to pause the workflow, 'resume' to resume a paused workflow."
        elif "status" in user_input.lower():
            response = f"Current workflow step: {state.get('current_step', 'Unknown')}"
        elif "pause" in user_input.lower():
            response = "Workflow paused. Type 'resume' to continue."
            return {"user_feedback": response, "current_step": "user_interaction_handler"}
        elif "resume" in user_input.lower():
            response = "Workflow resumed."
            return {"user_feedback": response, "current_step": state.get("last_successful_step", "workflow_controller_agent")}
        else:
            response = f"Processed user input: {user_input}. Type 'help' for available commands."


        end_time = time.time()
        elapsed_time = end_time - start_time
        token_usage = 0  # Assuming token usage is determined here

        state["agent_metrics"]["agent_name"].append("user_interaction_handler")
        state["agent_metrics"]["total_time"].append(elapsed_time)
        state["agent_metrics"]["total_tokens"].append(token_usage)
        return {"user_input": "", "user_feedback": response, "agent_metrics": state["agent_metrics"]}
    except Exception as e:
        logger.error(f"Error in user_interaction_handler: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return {"errors": [f"User interaction error: {str(e)}"], "context_request": "user_interaction_failed"}
		
def create_audit_workflow():
    workflow = StateGraph(State)

    # Add nodes
    workflow.add_node("workflow_controller_agent", workflow_controller_agent)
    workflow.add_node("preprocessor_agent", preprocessor_agent)
    workflow.add_node("risk_assessment_agent", risk_assessment_agent)
    workflow.add_node("report_generator_agent", report_generator_agent)
    workflow.add_node("error_handler", error_handler)
    workflow.add_node("user_interaction_handler", user_interaction_handler)

    # Define conditional edges
    workflow.add_conditional_edges(
        "workflow_controller_agent",
        lambda x: x["current_step"],
        {
            "preprocessor_agent": "preprocessor_agent",
            "risk_assessment_agent": "risk_assessment_agent",
            "report_generator_agent": "report_generator_agent",
            "error_handler": "error_handler",
            "user_interaction_handler": "user_interaction_handler",
            END: END
        }
    )

    # Add edges back to workflow_controller_agent
    for node in ["preprocessor_agent", "risk_assessment_agent", 
                 "report_generator_agent", "error_handler", "user_interaction_handler"]:
        workflow.add_edge(node, "workflow_controller_agent")

    # Set entry point
    workflow.set_entry_point("workflow_controller_agent")

    return workflow

# Compile the workflow
workflow = create_audit_workflow()
app = workflow.compile()

try:
    display(Image(app.get_graph(xray=True).draw_mermaid_png()))
except Exception:
    pass

def run_audit(initial_state: Dict[str, Any] = None):
    experiment_id = setup_mlflow_experiment()
    state = initial_state or {}
    state = State(**state)  # Ensure state conforms to State schema
    
    results = []
    
    with mlflow.start_run(experiment_id=experiment_id) as run:
        try:
            # Log input parameters
            mlflow.log_params({
                "workflow_start_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "num_input_documents": len(state.get("raw_documents", {})),
            })
            
            # Initialize metrics dictionary for accumulating metrics
            workflow_metrics = {
                "total_processing_time": 0,
                "total_token_usage": 0,
                "num_errors": 0,
                "num_steps_completed": 0
            }
            
            for step_output in app.stream(state):
                step_name, step_result = list(step_output.items())[0]
                logger.info(f"Executed: {step_name}")
                
                # Log step-specific metrics
                if "agent_metrics" in step_result:
                    metrics = step_result["agent_metrics"]
                    for i, agent_name in enumerate(metrics["agent_name"]):
                        mlflow.log_metrics({
                            f"{agent_name}_processing_time": metrics["total_time"][i],
                            f"{agent_name}_token_usage": metrics["total_tokens"][i]
                        })
                        
                        # Update accumulated metrics
                        workflow_metrics["total_processing_time"] += metrics["total_time"][i]
                        workflow_metrics["total_token_usage"] += metrics["total_tokens"][i]
                
                # Log any errors
                if "errors" in step_result and step_result["errors"]:
                    workflow_metrics["num_errors"] += len(step_result["errors"])
                    for i, error in enumerate(step_result["errors"]):
                        mlflow.log_text(str(error), f"errors/error_{workflow_metrics['num_errors']}_{i}.txt")
                
                workflow_metrics["num_steps_completed"] += 1
                
                # Update state and results
                results.append({"step": step_name, "result": step_result})
                state.update(step_result)
                
                if step_name == END:
                    logger.info("Audit workflow completed")
                    break
            
            # Log final accumulated metrics
            mlflow.log_metrics(workflow_metrics)
            
            # Log output artifacts
            if "draft_report" in state:
                mlflow.log_text(state["draft_report"], "draft_report.txt")
            
            if "docx_report" in state:
                mlflow.log_artifact(state["docx_report"])
            
            # Log final success/failure status
            mlflow.log_param("workflow_status", "completed" if "errors" not in state else "failed")
            
            # Log workflow completion time
            mlflow.log_param("workflow_end_time", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            
            # Log evaluation metrics if available
            if "risk_assessment" in state:
                risk_metrics = {
                    "num_high_risks": sum(1 for _, rating in state["risk_assessment"].get("assessment_rating", {}).items() if "high" in rating.lower()),
                    "num_medium_risks": sum(1 for _, rating in state["risk_assessment"].get("assessment_rating", {}).items() if "medium" in rating.lower()),
                    "num_low_risks": sum(1 for _, rating in state["risk_assessment"].get("assessment_rating", {}).items() if "low" in rating.lower())
                }
                mlflow.log_metrics(risk_metrics)
            
        except Exception as e:
            logger.error(f"Error in run_audit: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            mlflow.log_param("workflow_status", "failed")
            mlflow.log_text(str(e), "error.txt")
            mlflow.log_text(traceback.format_exc(), "error_traceback.txt")
            results.append({"step": "error", "result": {"errors": [str(e)]}})
    
    return results, state
    


w = WorkspaceClient()
dbutils = w.dbutils

def save_to_dbfs(file_content, file_path, filename):
    # Create the full DBFS path
    dbfs_path = file_path + filename
    # Base64 encode the binary file content (as DBFS needs string content)
    #file_content_encoded = base64.b64encode(file_content).decode('utf-8')
    # Write the content to DBFS using dbutils.fs.put
    w.files.upload(dbfs_path, file_content, overwrite=True)
    # Return the DBFS path (for tracking or further use)
    return dbfs_path

def gradio_interface(risk_matrix, control_testing, previous_audits):
    try:
        with mlflow.start_run(run_name="gradio_interface_run") as run:
            # Log input file information
            mlflow.log_params({
                "risk_matrix_size": len(risk_matrix),
                "control_testing_size": len(control_testing),
                "previous_audits_size": len(previous_audits)
            })
            
            # Save uploaded files and run audit
            save_to_dbfs(risk_matrix, "/Volumes/genaudit_catalog/genaudit_schema/genaudit_volume/", "risk_matrix.xlsx")
            save_to_dbfs(control_testing, "/Volumes/genaudit_catalog/genaudit_schema/genaudit_volume/", "control_testing.xlsx")
            save_to_dbfs(previous_audits, "/Volumes/genaudit_catalog/genaudit_schema/genaudit_volume/", "previous_audits.docx")

            risk_matrix_path = w.dbfs.read("dbfs:/Volumes/genaudit_catalog/genaudit_schema/genaudit_volume/risk_matrix.xlsx")
            control_testing_path = w.dbfs.read("dbfs:/Volumes/genaudit_catalog/genaudit_schema/genaudit_volume/control_testing.xlsx")
            previous_audits_path = w.dbfs.read("/Volumes/genaudit_catalog/genaudit_schema/genaudit_volume/previous_audits.docx")

            initial_state = {
                "current_step": "workflow_controller_agent",
                "raw_documents": {
                    "risk_matrix": risk_matrix_path,
                    "control_testing": control_testing_path,
                    "previous_audits": previous_audits_path
                },
                "preprocessed_data": {},
                "risk_assessment": {},
                "draft_report": "",
                "final_report": "",
                "context_request": "",
                "errors": [],
                "user_input": "",
                "user_feedback": "",
                "racm_data": [],
                "agent_metrics": {
                    "agent_name": [],
                    "total_time": [],
                    "total_tokens": []
                }
            }
            
            results, final_state = run_audit(initial_state)
            
            # Prepare output
            output = "Audit workflow completed successfully.\n\nResults:\n"
            for result in results:
                output += f"Step: {result['step']}\n"
                output += f"Result: {result['result']}\n"
                output += "---\n"
            
            # Log the final output
            mlflow.log_text(output, "output.txt")
            
            return output
            
    except Exception as e:
        error_msg = f"An error occurred during the audit workflow:\n\n{str(e)}\n\nFull traceback:\n{traceback.format_exc()}"
        mlflow.log_param("error", str(e))
        mlflow.log_text(error_msg, "error_details.txt")
        return error_msg

# Set up the Gradio interface
iface = gr.Interface(
    fn=gradio_interface,
    inputs=[
    gr.File(label="Risk Matrix (Excel file)", type="binary"),
    gr.File(label="Control Testing (Excel file)", type="binary"),
    gr.File(label="Previous Audits (Word document)", type="binary")
    ],
    outputs=gr.Textbox(label="Output", lines=20),
    title="Audit Workflow",
    description="Upload the required files to run the audit workflow."
)

# Run the Gradio app
iface.launch(debug=True)

# COMMAND ----------

# MAGIC %md
# MAGIC # Evaluation

# COMMAND ----------

def eval_model(risk_matrix, control_testing, previous_audits,eval_judge,final_state):
    # Define the paths to your evaluation data
    risk_matrix_path = risk_matrix
    #"dbfs:/Volumes/genaudit/genaudit/volume/Risk_Matrix_Excel.xlsx"
    control_testing_path = control_testing
    #"dbfs:/Volumes/genaudit/genaudit/volume/Accounts_Payable_Control_Testing_fixed_v1.xlsx"
    previous_audits_path = previous_audits
    #"/Volumes/genaudit/genaudit/volume/Procurement_Audit_Report_2023_Synthetic.docx"
    #generated_report_path = "/Volumes/labuser7921563_1728624720_t7ow_da/genaudit_v2/genaudit_v2/Audit_Report_v7.docx"

    # Load the risk matrix data
    risk_matrix_df = ps.read_excel(risk_matrix_path, sheet_name='Sheet1')
    risk_matrix = risk_matrix_df.to_string(index=False)

    # Load the control testing data and only those with No adequacy
    control_testing_df = ps.read_excel(control_testing_path, sheet_name='RACM')
    filtered_racm_data = control_testing_df[control_testing_df["Control Adequacy (Y/N)"] == "No"] 
    control_testing_data = control_testing_df.to_string(index=False)



    # Content Accuracy examples, rubrics, and metric creation
    accuracy_example_score_1 = EvaluationExample(
        input="Control testing data for an internal audit on network security controls. The data includes various sections like access control, incident response, and configuration management.",
        output=(
        
            " Access Control: High Risk"
            "  + The company has a robust access control system in place, with proper segregation of duties and least privilege access."
            "  + However, there are some areas of improvement identified, including:"
            "    - Some users have excessive access rights. "
            "    - Some access control rules are not properly documented. "
            "* Incident Response: High Risk"
            "  + The company has an incident response plan in place, with clear procedures for responding to security incidents."
            "  + However, there are some areas of improvement identified, including: "
            "    - The incident response plan is not properly tested."
            "    - Some incident response procedures are not properly documented."
            "* Configuration Management: Medium Risk"
            "  + The company has a configuration management system in place, with proper change management and version control."
            "  + However, there are some areas of improvement identified, including:"
            "    - Some configuration changes are not properly documented."
            "    - Some configuration settings are not properly secured."
        
        ),
        score=3,
        justification=(
                    
            "The report accurately reflects the control testing data and risk matrix provided. "
            "The report correctly identifies the controls in place for access control, incident response, and configuration management."
            "The report also accurately identifies the areas of improvement for each control, including excessive access rights, lack of documentation, and untested incident response plans."
            "The report provides a clear and concise summary of the control testing data and risk matrix, and to an extent reflects the level of risk associated with each control."
        ),
    )
    accuracy_example_score_2 = EvaluationExample(
        input="Control testing data for an internal audit on network security controls. The data includes various sections like access control, incident response, and configuration management.",
        output=(
                
            "Access Control: The company has a robust access control system in place. However, there are no areas of improvement identified. "
            "Incident Response: The company has an incident response plan in place. However, there are no areas of improvement identified. "
            "Configuration Management: The company has a configuration management system in place. However, there are no areas of improvement identified."
        

        ),
        score=1,
        justification=(
                "The report fails to accurately reflect the control testing data and risk matrix provided. "
            "The report does not identify any areas of improvement for access control, incident response, or configuration management, despite the presence of several high-risk findings in the control testing data. "
            "The report provides a cursory summary of the control testing data and risk matrix, but fails to provide any meaningful analysis or recommendations for improvement. "
            "The report's failure to accurately reflect the control testing data and risk matrix, and its lack of meaningful analysis and recommendations, make it a low-quality report."
        
        ),
    )   




    eval_rubrics_accuracy = [
        {
            "metric": "fact_checking",
            "rubrics": """
            Score 1: The audit report contains significant factual inaccuracies or incorrect details that contradict the provided control testing data,  or risk matrix.
            Score 2: The audit report contains some factual inaccuracies, but the core findings are mostly aligned with the provided input data.
            Score 3: The audit report is factually correct, with only minor issues or missing details that do not significantly affect the overall accuracy.
            Score 4: The audit report is completely accurate, with all findings, data points, and conclusions fully supported by the provided control testing data,  and risk matrix.
            """,
        },
        {
            "metric": "content_accuracy",
            "rubrics": """
            Score 1: The content of the audit report does not accurately represent the control testing or risk matrix, and there are substantial deviations.
            Score 2: The content of the audit report has some inaccuracies or deviations, but it captures the main points of the input data.
            Score 3: The content of the audit report is mostly accurate, with only minor deviations or missing information.
            Score 4: The content of the audit report is fully accurate, with no deviations or missing information. It provides a clear and precise reflection of the control testing, and risk matrix.
            """,
        },
    ]




    prompt_accuracy = f"""
    You are an expert auditor tasked with evaluating the accuracy and fact-checking of the an audit report, based on the control testing data and risk matrix provided. Please assess the report on the following metrics:

    1. **Fact-Checking**: How well does the audit report reflect the actual data from the control testing, and risk matrix? Are all the facts correct?
    2. **Content Accuracy**: Does the audit report accurately represent the key findings and conclusions based on the provided input? Are there any deviations or missing information?

    Control Testing Data: {control_testing_data}

    Risk Matrix: {risk_matrix}

    please provide a score from 1 to 4 for each metric based on the rubrics below and include a brief explanation for each score.

    ### Fact-Checking Rubric:
    {eval_rubrics_accuracy[0]['rubrics']}

    ### Content Accuracy Rubric:
    {eval_rubrics_accuracy[1]['rubrics']}
    """

    content_accuracy = mlflow.metrics.genai.make_genai_metric(
        name="Accuracy",
        definition=(
            "Completeness refers to whether the audit report covers all necessary sections, key findings, and recommendations "
            "based on the provided control testing data . An incomplete report may omit "
            "important details, sections, or findings that are critical to understanding the full context of the audit."
        ),
        grading_prompt= (prompt_accuracy),
        examples=[
            accuracy_example_score_1, 
            accuracy_example_score_2
        ],
        model= eval_judge,
        parameters={"temperature": 0.0},
        aggregations=["mean", "variance"],
        greater_is_better=True
        
    )

    # Structure Accuracy examples, rubrics, and metric creation
    structure_example_score_1 = EvaluationExample(
        input="Control testing data for an internal audit on network security controls.",
        output=(
        
        "Audit Report: Introduction: This audit report covers the control testing data for the internal audit on network security controls. "
            "Executive Summary: The audit report provides an overview of the control testing data and risk matrix. "
            "Methodology: The audit was conducted using a risk-based approach, with a focus on identifying high-risk areas. "
            "Key Findings: Access Control: The company has a robust access control system in place. However, there are some areas of improvement identified. "
            "Incident Response: The company has an incident response plan in place. However, there are some areas of improvement identified. "
            "Configuration Management: The company has a configuration management system in place. However, there are some areas of improvement identified. "
            "Recommendations: The report provides recommendations for improving the control environment, including implementing additional controls and enhancing existing controls. "
            "Conclusion: The audit report concludes that the company has a robust control environment in place, but there are some areas of improvement identified."
    
        
        ),
        score=4,
        justification=(
        "The report has a clear and logical structure, with an introduction, executive summary, methodology, key findings, recommendations, and conclusion. "
            "The report follows a standard audit report format, making it easy to follow and understand. "
            "The report's structure and organization make it easy to navigate and understand."
        
        ),
    )
    structure_example_score_2 = EvaluationExample(
        input="Control testing data for an internal audit on network security controls.",
        output=(
                
            "Audit Report: This audit report covers the control testing data for the internal audit on network security controls. "
            "Some findings were identified, including access control issues and incident response plan weaknesses. "
            "Recommendations are provided to improve the control environment, but they are not clearly linked to the findings. "
            "There is no clear methodology or approach described in the report, and the language is often unclear and confusing."
    
        ),
        score=1,
        justification=(
                "The report lacks a clear and logical structure, with no introduction, executive summary, or conclusion. "
            "The report does not follow a standard audit report format, making it difficult to follow and understand. "
            "The report's language is often unclear and confusing, making it difficult to navigate and understand."   
        
        ),
    )   



    eval_rubrics_structure = [
        {
            "metric": "structure_accuracy",
            "rubrics": """
            Score 1: The report lacks a clear and logical structure, with missing or disorganized sections. The stakeholder list is incomplete or inaccurate. The executive summary is missing or does not provide a clear overview of the audit. The summary of audit issues and risk ratings is missing or incomplete. The key findings and management actions are not clearly presented. The detailed findings section is missing or does not provide sufficient information. The conclusion is missing or does not summarize the main points of the report.
            
            Score 2:The report has some structural issues, but the main sections are present. The stakeholder list is mostly accurate, but may be missing some information. The executive summary provides a brief overview of the audit, but may not be clear or concise. The summary of audit issues and risk ratings is mostly complete, but may be missing some information. The key findings and management actions are presented, but may not be clearly organized. The detailed findings section provides some information, but may be incomplete or disorganized. The conclusion summarizes the main points of the report, but may not be clear or concise.
            
            Score 3:The report has a clear and logical structure, with all main sections present. The stakeholder list is accurate and complete. The executive summary provides a clear and concise overview of the audit. The summary of audit issues and risk ratings is complete and accurate. The key findings and management actions are clearly presented and organized. The detailed findings section provides sufficient information and is well-organized. The conclusion summarizes the main points of the report clearly and concisely.

            Score 4: The report has a clear and logical structure, with all main sections present and well-organized. The stakeholder list is accurate and complete, with clear and concise information. The executive summary provides a clear and concise overview of the audit, with all necessary information. The summary of audit issues and risk ratings is complete and accurate, with clear and concise information. The key findings and management actions are clearly presented and organized, with all necessary information. The detailed findings section provides sufficient information and is well-organized, with clear and concise language. The conclusion summarizes the main points of the report clearly and concisely, with all necessary information.
            """,
        },
    ]



    prompt_structure = f"""
    You are an expert auditor tasked with evaluating the structure and organization of an audit report. Please assess the report on the following metric:

    1. **Structure Accuracy**: How well is the audit report structured and organized? Does it follow the structure guidelines of title, stakeholder list, for Action, for information, audit team, executive summary, audit entity, report date, summary of audit issues and risk ratings, key findings and management actions, detailed findings, issue category, observation, justification, details, recommendation and conclusion)?

    Please provide a score from 1 to 4 for this metric based on the rubric below and include a brief explanation for the score.

    ### Structure Accuracy Rubric:
    {eval_rubrics_structure[0]['rubrics']}
    """




    structure_accuracy = mlflow.metrics.genai.make_genai_metric(
        name="Structure",
        definition=(
            "Structure Accuracy refers to whether the audit report is well-organized, logical, and easy to follow. "
            "A well-structured report should have a clear introduction, executive summary, and conclusion, "
            "and should present the key findings and recommendations in a clear and concise manner."
        ),
        grading_prompt= (prompt_structure),
        examples=[
            structure_example_score_1, 
            structure_example_score_2
        ],
        model= eval_judge,
        parameters={"temperature": 0.0},
        aggregations=["mean", "variance"],
        greater_is_better=True
        
    )

    # Readability Evaluation  examples, rubrics, and metric creation
    readability_example_score_1 = EvaluationExample(
        input="Evaluate the readability of the following audit report.",
        output="We also recomend that the company hire a new accountant who can help they're with they're accounting and financial reporting." 
            "This will help the company to improve they're financial management and reduce they're risk."

            "In conclusion, the company have a lot of problem with they're accounting and financial reporting. "
            "They need to take step to improve they're financial management and reduce they're risk. "
            "We recomend that the company implement a new accounting system, develop a plan for reporting they're financial information, and hire a new accountant.",
        score=1,
        justification="The report is hard to read due to frequent grammar mistakes and a lack of clear flow. It does not meet professional standards."
    )
    
    readability_example_score_2 = EvaluationExample(
        input="Evaluate the readability of the following audit report.",
        output="We recommend that the company take the following steps to improve its financial reporting and risk management practices:"

            "* Develop a more comprehensive risk management plan"
            "* Regularly review and update the risk management plan"
            "* Improve the categorization of expenses in the financial statements"
            "* Provide more detailed information about the company's financial performance"

            "In conclusion, the company's financial reporting and risk management practices are generally sound, but there are some areas for improvement. We recommend that the company take the steps outlined above to improve its financial reporting and risk management practices.",
        score=3,
        justification="The report is generally well-written, with some areas that could benefit from improved clarity and professionalism."
    )


    readability_rubrics = [ 
        {
            "metric": "readability",
            "rubrics": """
    Score 1: The report is unclear, poorly structured, and contains significant grammatical or stylistic issues. The professionalism is lacking, and it is difficult to follow.
    Score 2: The report has some clarity but contains noticeable issues with coherence, grammar, or professionalism. Some sections may be confusing or hard to understand.
    Score 3: The report is mostly clear and coherent, with only minor grammar or professionalism issues. It reads well but could be improved in certain areas.
    Score 4: The report is highly clear, well-organized, and free of grammar or professionalism issues. It is easy to read and follows a logical flow, demonstrating strong professionalism.
    """,
    },
                    ]
 

    readability_prompt = f"""
    You are an expert auditor tasked with evaluating the readability of an audit report based on clarity, coherence, grammar, and professionalism. Please assess the report on the following metrics:
    
    1. **Clarity**: Is the language clear and easy to understand?
    2. **Coherence**: Does the report flow logically from one section to the next?
    3. **Grammar**: Are there any grammatical errors or awkward phrasing?
    4. **Professionalism**: Does the report maintain a formal and professional tone throughout?
    
    Please provide a score from 1 to 4 for each metric based on the rubric below and include a brief explanation for each score.
    
    ### Readability Rubric:
    {readability_rubrics[0]['rubrics']}
    """

    readability = mlflow.metrics.genai.make_genai_metric(
        name="Readability",
        definition=(
            "Readability refers to how easy the audit report is to read and understand, including clarity, coherence, grammar, and professionalism."
        ),
        grading_prompt=readability_prompt,
        examples=[readability_example_score_1, readability_example_score_2],
        model= eval_judge,
        parameters={"temperature": 0.0},
        aggregations=["mean", "variance"],
        greater_is_better=True
    )

    # prepare the evaluation data and use the metrics to evaluate
    eval_data = pd.DataFrame([{"draft_report": final_state.get("draft_report")}])
    # The first argument of mlflow.evaluate needs to be either an endpoint or function
    def genaudit_model(df):
        content_parts = df["draft_report"]
    return content_parts
    # Evaluate the model
    eval_result = mlflow.evaluate(
    model=genaudit_model,
    data=eval_data,
    model_type="question-answering",
    extra_metrics=[content_accuracy, structure_accuracy, readability],
    evaluator_config={'default': {'col_mapping': {'inputs': 'draft_report'}}}
    )

    #Combine the model evaluation with the agent evaluation
    eval_df = eval_result.tables["eval_results_table"]
    df_agent_metrics = pd.DataFrame(final_state["agent_metrics"])
    # Summing up the numbers for each agent_name
    summed_metrics = df_agent_metrics.groupby('agent_name').sum()

    # Resetting index to make 'agent_name' a column again if needed
    summed_metrics.reset_index(inplace=True)
    agent_metrics = pd.DataFrame([summed_metrics.iloc[:,0:].to_dict()], index=None)

    # Concatenate summed_metrics as a new record to eval_result.tables["eval_results_table"]
    eval_model = pd.concat(
        [eval_result.tables["eval_results_table"], agent_metrics],
        axis=1
        
    )


    # Set display options to show the full content of the column
    pd.set_option('display.max_colwidth', None)


    # Display the updated eval_results_table to verify the addition
    display(eval_model[['toxicity/v1/score',
        'flesch_kincaid_grade_level/v1/score', 'ari_grade_level/v1/score',
        'Accuracy/v1/score', 'Accuracy/v1/justification', 'Structure/v1/score',
        'Structure/v1/justification', 'Readability/v1/score',
        'Readability/v1/justification', 'agent_name', 'total_time',
        'total_tokens']])


